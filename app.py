import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.cluster import KMeans, DBSCAN, AgglomerativeClustering
from sklearn.mixture import GaussianMixture
from sklearn.metrics import silhouette_score, davies_bouldin_score, calinski_harabasz_score
from sklearn.decomposition import PCA
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
import io
import os
import warnings
from PIL import Image # For loading image

# Suppress warnings for cleaner app output
warnings.filterwarnings("ignore")

# --- Session State Initialization ---
if "analysis_completed" not in st.session_state:
    st.session_state.analysis_completed = False
if "file_uploader_key" not in st.session_state:
    st.session_state.file_uploader_key = 0
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

# Page config
st.set_page_config(
    layout="wide",
    page_title="Unsupervised Learning (using Machine Learning)" # Updated title
)

# --- Authentication Logic ---
def login_page():
    # Attempt to load logo for login page
    try:
        logo = Image.open("Logo.png")
        st.image(logo, width=150) # Adjust width as needed
    except FileNotFoundError:
        st.warning("Logo.png not found. Please ensure it's in the repository.")

    st.title("🔒 Login to Unsupervised Learning Dashboard") # Updated login title
    st.markdown("Please enter your credentials to access the application.")

    username_input = st.text_input("Username")
    password_input = st.text_input("Password", type="password")

    if st.button("Login"):
        try:
            # Retrieve the list of valid users from secrets
            # st.secrets["users"] is expected to be a list of dictionaries like:
            # [
            #     {"username": "user1", "password": "password1"},
            #     {"username": "user2", "password": "password2"}
            # ]
            valid_users = st.secrets["users"]
        except KeyError:
            st.error("Secrets not configured correctly! Please set 'users' as a list of username/password dictionaries in Streamlit Cloud secrets or secrets.toml.")
            return

        authenticated = False
        for user_data in valid_users:
            if username_input == user_data["username"] and password_input == user_data["password"]:
                authenticated = True
                break # Exit loop once a match is found

        if authenticated:
            st.session_state.authenticated = True
            st.success("Login successful! Redirecting...")
            st.rerun() # Rerun to switch to the main app content
        else:
            st.error("Invalid username or password.")

    # --- Copyright Notice for Login Page ---
    st.markdown("---") # Optional: add a separator line
    st.markdown("<p style='text-align: center; color: grey;'>© Copyright SSO Consultants</p>", unsafe_allow_html=True)


# --- Main Application Content (rest of your app.py code remains the same) ---
def main_app():
    # ... (your existing main_app code goes here, from the logo in sidebar onwards)
    # This part of the code is unchanged from your previous app (20).py file.
    # It starts with:
    # st.sidebar.title(" ")
    # try:
    #     logo = Image.open("Logo.png")
    #     st.sidebar.image(logo, use_container_width=True)
    # ... and continues to the end of the file.

    st.sidebar.title(" ") # Small space before logo
    try:
        logo = Image.open("Logo.png")
        # FIX: Replace use_column_width with use_container_width
        st.sidebar.image(logo, use_container_width=True) # Use container_width for responsiveness
    except FileNotFoundError:
        st.sidebar.warning("Logo.png not found. Please ensure it's in the repository.")

    st.title("📊 Unsupervised Learning (using Machine Learning)") # Main heading updated

    st.markdown("""
    Welcome! This app helps you discover customer segments using unsupervised machine learning.
    """)

    # --- Logout Button ---
    if st.session_state.authenticated:
        if st.sidebar.button("Logout"):
            st.session_state.authenticated = False
            st.rerun()

    # --- Helper Functions (defined inside main_app or globally if @st.cache_data is used) ---

    # Function to safely format metrics or return 'N/A'
    def format_metric(value):
        return f'{value:.4f}' if isinstance(value, (int, float)) else "N/A"

    # Function to detect potential ID columns
    def detect_potential_id_columns(df, uniqueness_threshold=0.9):
        potential_ids = []
        for col in df.columns:
            # Check for common ID-like names
            if any(keyword in col.lower() for keyword in ['id', 'user_id', 'customer_id', 'client_id']):
                potential_ids.append(col)
                continue
            
            # Check uniqueness ratio for non-numeric or string-like numeric columns
            temp_col = df[col].copy()
            try:
                # Try to convert to numeric, coercing errors
                numeric_check = pd.to_numeric(temp_col, errors='coerce')
                # If it's mostly numbers, but still has high unique count, it might be an ID
                if not numeric_check.isnull().all(): # Make sure it's not all NaNs after coercion
                    if numeric_check.nunique() / len(df) > uniqueness_threshold:
                        potential_ids.append(col)
                        continue
            except Exception:
                # If not numeric, check uniqueness for string/object types
                if df[col].dtype == 'object' and df[col].nunique() / len(df) > uniqueness_threshold:
                    potential_ids.append(col)
                    continue
        return list(set(potential_ids)) # Use set to remove duplicates

    @st.cache_data
    def preprocess_data(df, numeric, categorical, missing):
        # Ensure only selected features are passed for preprocessing
        df_proc = df[numeric + categorical].copy()
        
        if missing == "drop_rows":
            df_proc.dropna(inplace=True)
        else:
            for col in numeric:
                df_proc[col].fillna(df_proc[col].mean(), inplace=True)
            for col in categorical:
                df_proc[col].fillna(df_proc[col].mode()[0], inplace=True)
        
        # Store df_for_profile before one-hot encoding for accurate profile generation
        # It should include original numeric and categorical columns
        df_for_profile = df_proc.copy()

        encoded_features = []
        if categorical:
            encoder = OneHotEncoder(sparse_output=False, handle_unknown="ignore")
            enc_data = encoder.fit_transform(df_proc[categorical])
            enc_df = pd.DataFrame(enc_data, columns=encoder.get_feature_names_out(categorical), index=df_proc.index)
            df_proc.drop(columns=categorical, inplace=True)
            df_proc = pd.concat([df_proc, enc_df], axis=1)
            encoded_features = encoder.get_feature_names_out(categorical).tolist()

        scaler = StandardScaler()
        scaled = scaler.fit_transform(df_proc)
        scaled_df = pd.DataFrame(scaled, columns=df_proc.columns, index=df_proc.index)
        return scaled_df, df_for_profile

    @st.cache_data
    def evaluate_models(scaled_df, k_range):
        scores = {}
        inertia_kmeans = []
        aic_gmm = []
        bic_gmm = []

        for algo in ["KMeans", "GMM", "Agglomerative"]:
            scores[algo] = {"Silhouette": [], "Davies": [], "Calinski": []}
            for k in k_range:
                if k >= len(scaled_df):
                    scores[algo]["Silhouette"].append(np.nan)
                    scores[algo]["Davies"].append(np.nan)
                    scores[algo]["Calinski"].append(np.nan)
                    if algo == "KMeans":
                        inertia_kmeans.append(np.nan)
                    elif algo == "GMM":
                        aic_gmm.append(np.nan)
                        bic_gmm.append(np.nan)
                    continue
                try:
                    if algo == "KMeans":
                        model = KMeans(n_clusters=k, n_init=10, random_state=42)
                        labels = model.fit_predict(scaled_df)
                    elif algo == "GMM":
                        model = GaussianMixture(n_components=k, random_state=42)
                        labels = model.fit_predict(scaled_df)
                        aic_gmm.append(model.aic(scaled_df))
                        bic_gmm.append(model.bic(scaled_df))
                    else: # Agglomerative
                        labels = AgglomerativeClustering(n_clusters=k).fit_predict(scaled_df)

                    if len(np.unique(labels)) > 1:
                        scores[algo]["Silhouette"].append(silhouette_score(scaled_df, labels))
                        scores[algo]["Davies"].append(davies_bouldin_score(scaled_df, labels))
                        scores[algo]["Calinski"].append(calinski_harabasz_score(scaled_df, labels))
                    else:
                        scores[algo]["Silhouette"].append(np.nan)
                        scores[algo]["Davies"].append(np.nan)
                        scores[algo]["Calinski"].append(np.nan)
                except Exception:
                    scores[algo]["Silhouette"].append(np.nan)
                    scores[algo]["Davies"].append(np.nan)
                    scores[algo]["Calinski"].append(np.nan)
                    if algo == "KMeans":
                        inertia_kmeans.append(np.nan)
                    elif algo == "GMM":
                        aic_gmm.append(np.nan)
                        bic_gmm.append(np.nan)
        
        return scores, {"KMeans_Inertia": inertia_kmeans, "GMM_AIC": aic_gmm, "GMM_BIC": bic_gmm}

    @st.cache_data
    def generate_plots(df, labels, numeric_cols, pca_components=2):
        """Generates PCA plot and cluster profile plots."""
        if labels is None or len(np.unique(labels)) <= 1:
            return None, None, None, None # Return Nones for all outputs if conditions not met
        
        df_clustered = df.copy()
        # IMPORTANT FIX: Do NOT convert labels to string here. Keep as int for correct comparison with actual labels.
        df_clustered['Cluster'] = labels # Keep as int for correct size calculation later

        # PCA Plot
        # Ensure there are enough features for PCA
        if len(numeric_cols) >= pca_components:
            pca = PCA(n_components=min(pca_components, len(numeric_cols)), random_state=42)
            # Use only numeric columns for PCA
            components = pca.fit_transform(df_clustered[numeric_cols])
            pca_df = pd.DataFrame(data=components, columns=[f'PC{i+1}' for i in range(components.shape[1])])
            pca_df['Cluster'] = df_clustered['Cluster'] # Add numeric cluster labels
            
            # Convert 'Cluster' column to object/string only for hue in seaborn for categorical color mapping
            pca_df['Cluster_Str'] = pca_df['Cluster'].astype(str)

            fig_pca, ax_pca = plt.subplots(figsize=(10, 7))
            sns.scatterplot(x='PC1', y='PC2', hue='Cluster_Str', data=pca_df, palette='viridis', ax=ax_pca, s=100, alpha=0.7)
            ax_pca.set_title('PCA of Clusters')
            ax_pca.set_xlabel(f'Principal Component 1 ({pca.explained_variance_ratio_[0]*100:.2f}%)')
            ax_pca.set_ylabel(f'Principal Component 2 ({pca.explained_variance_ratio_[1]*100:.2f}%)')
            ax_pca.legend(title='Cluster')
            plt.close(fig_pca)
        else:
            fig_pca = None # Set to None if PCA cannot be performed

        # Cluster Profiles (Mean for Numeric, Proportions for Categorical)
        fig_profile_numeric, ax_profile_numeric = plt.subplots(figsize=(12, 6))
        if numeric_cols: # Only plot if numeric columns exist
            cluster_means_numeric = df_clustered.groupby('Cluster')[numeric_cols].mean()
            cluster_means_numeric.T.plot(kind='bar', ax=ax_profile_numeric, colormap='viridis')
            ax_profile_numeric.set_title('Numeric Feature Means by Cluster')
            ax_profile_numeric.set_ylabel('Mean Value')
            ax_profile_numeric.tick_params(axis='x', rotation=45)
            ax_profile_numeric.legend(title='Cluster')
            plt.tight_layout()
        else:
            fig_profile_numeric = None # Set to None if no numeric columns
        plt.close(fig_profile_numeric)

        # Categorical Feature Proportions
        cluster_cat_proportions = {}
        # Get actual categorical columns from df_clustered, excluding the 'Cluster' column itself
        original_cat_in_df_clustered = [col for col in df_clustered.columns if df_clustered[col].dtype == 'object' and col != 'Cluster']
        
        for col in original_cat_in_df_clustered:
            if col in df_clustered.columns: 
                proportions = df_clustered.groupby('Cluster')[col].value_counts(normalize=True).unstack(fill_value=0)
                cluster_cat_proportions[col] = proportions

        return fig_pca, fig_profile_numeric, cluster_means_numeric, cluster_cat_proportions

    # Modified generate_cluster_summaries to return structured data
    def generate_cluster_summaries(cluster_means_numeric, cluster_cat_proportions, original_df_for_profile, labels, chosen_algo):
        structured_summaries = {}
        total_samples = len(original_df_for_profile) # This is the original *processed* df size

        # Calculate overall means and stds for numeric features
        overall_means = None
        overall_stds = None
        if cluster_means_numeric is not None and not cluster_means_numeric.empty:
            # Ensure we're only looking at numeric columns that were actually clustered
            relevant_numeric_cols = [col for col in cluster_means_numeric.columns if col in original_df_for_profile.columns]
            if relevant_numeric_cols:
                overall_means = original_df_for_profile[relevant_numeric_cols].mean()
                overall_stds = original_df_for_profile[relevant_numeric_cols].std()

        # Get counts for each cluster directly from the labels array
        unique_labels, counts = np.unique(labels, return_counts=True)
        label_counts = dict(zip(unique_labels, counts))

        # Get all unique cluster IDs from the labels, including -1 for DBSCAN if present
        all_cluster_ids = sorted([c_id for c_id in np.unique(labels) if pd.notna(c_id)]) # Filter out potential NaNs if any

        for cluster_id_int in all_cluster_ids:
            cluster_size = label_counts.get(cluster_id_int, 0)

            # Special handling for DBSCAN's noise cluster (-1)
            if chosen_algo == "DBSCAN" and cluster_id_int == -1:
                 if cluster_size == 0: continue # If no noise points, skip
                 structured_summaries[str(cluster_id_int)] = { # Use string key for consistency with potentially mixed dict keys later
                     "cluster_heading": f"Noise Points (DBSCAN): (N={cluster_size}, {(cluster_size / total_samples) * 100:.1f}% of total data)",
                     "numeric_characteristics": [],
                     "categorical_characteristics": [],
                     "persona_implications": "These points could not be assigned to any distinct cluster based on the DBSCAN parameters (eps and min_samples). Potential Implications: *[These might be outliers, or your DBSCAN parameters might need adjustment to capture more clusters.]*"
                 }
                 continue # Move to next cluster_id

            cluster_percentage = (cluster_size / total_samples) * 100

            numeric_descriptors = []
            # Check if cluster_id_int is in the index of cluster_means_numeric before accessing
            if cluster_means_numeric is not None and cluster_id_int in cluster_means_numeric.index:
                for col in cluster_means_numeric.columns:
                    cluster_mean = cluster_means_numeric.loc[cluster_id_int, col] # Access by integer index

                    if overall_means is not None and col in overall_means and col in overall_stds:
                        overall_mean = overall_means[col]
                        overall_std = overall_stds[col]

                        if overall_std > 1e-9: # Avoid division by zero for std dev
                            z_score = (cluster_mean - overall_mean) / overall_std

                            if z_score > 1.2: # Stricter threshold for 'significantly higher'
                                numeric_descriptors.append(f"Higher {col.replace('_', ' ').title()}: Averaging {cluster_mean:.2f} (significantly above overall average of {overall_mean:.2f}).")
                            elif z_score < -1.2: # Stricter threshold for 'significantly lower'
                                numeric_descriptors.append(f"Lower {col.replace('_', ' ').title()}: Averaging {cluster_mean:.2f} (significantly below overall average of {overall_mean:.2f}).")
                        elif cluster_mean > overall_mean: # If std is 0, but cluster mean is higher
                            numeric_descriptors.append(f"Higher {col.replace('_', ' ').title()}: Averaging {cluster_mean:.2f}.")
                        elif cluster_mean < overall_mean: # If std is 0, but cluster mean is lower
                            numeric_descriptors.append(f"Lower {col.replace('_', ' ').title()}: Averaging {cluster_mean:.2f}.")

            categorical_descriptors = []
            for cat_col, proportions_df in cluster_cat_proportions.items():
                if cluster_id_int in proportions_df.index: # Access by integer index
                    cluster_proportions = proportions_df.loc[cluster_id_int].sort_values(ascending=False)

                    for category, cluster_prop in cluster_proportions.items():
                        if cluster_prop == 0: continue # Skip categories with 0 proportion
                        overall_prop = original_df_for_profile[cat_col].value_counts(normalize=True).get(category, 0)

                        if cluster_prop > 0.3 and cluster_prop > overall_prop * 1.5:
                            categorical_descriptors.append(f"Predominantly {category} (for {cat_col.replace('_', ' ').title()}): {cluster_prop:.1%} of this cluster falls into this category, which is significantly higher than the overall average ({overall_prop:.1%}).")
                        elif cluster_prop > 0.6:
                             categorical_descriptors.append(f"Majorly {category} (for {cat_col.replace('_', ' ').title()}): This category constitutes {cluster_prop:.1%} of the cluster.")

            persona_text = "Potential Persona/Implications: *[Consider giving this cluster a descriptive name like 'High-Value Customers' or 'New Engagers' based on its characteristics. Think about what these features mean for your business strategies.]*"

            if not numeric_descriptors and not categorical_descriptors:
                persona_text = "This cluster does not show strong deviations from the overall average in the selected features and may represent an 'average' segment. " + persona_text

            structured_summaries[str(cluster_id_int)] = { # Use string key for the output dictionary
                "cluster_heading": f"Cluster {cluster_id_int}: (N={cluster_size}, {cluster_percentage:.1f}% of total data)",
                "numeric_characteristics": numeric_descriptors,
                "categorical_characteristics": categorical_descriptors,
                "persona_implications": persona_text
            }
        return structured_summaries


    # Re-integrated create_report function
    def create_report(document, algorithm, params, metrics, data_preview_df, pca_plot_bytes, profile_plot_bytes, cluster_means_numeric, cluster_cat_proportions, original_df_for_profile, labels, chosen_algo):
        """Generates a comprehensive Word document report."""
        
        # --- Add logo to report ---
        try:
            logo = Image.open("Logo.png")
            logo_stream = io.BytesIO()
            logo.save(logo_stream, format="PNG")
            logo_stream.seek(0)
            document.add_picture(logo_stream, width=Inches(1.5)) # Adjust width as needed
            last_paragraph = document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center the logo
        except FileNotFoundError:
            document.add_paragraph("Logo (Logo.png) not found for report.")
        
        document.add_heading('ML Analysis Report', level=1) # Report title updated
        document.add_paragraph(f"Report generated on: {pd.to_datetime('today').strftime('%Y-%m-%d %H:%M:%S')}")

        document.add_heading('1. Analysis Overview', level=2)
        document.add_paragraph(
            "This report details the unsupervised learning analysis performed using the Streamlit application. "
            "The goal is to group similar data points based on their attributes, enabling targeted strategies."
        )

        document.add_heading('2. Clustering Parameters', level=2)
        document.add_paragraph(f"Algorithm Used: {algorithm}")
        for param, value in params.items():
            if value is not None: # Only add if parameter has a value
                document.add_paragraph(f"- {param.replace('_', ' ').title()}: {value}")
        
        document.add_heading('3. Model Performance Metrics', level=2)
        silhouette_str = format_metric(metrics.get('silhouette'))
        davies_bouldin_str = format_metric(metrics.get('davies_bouldin'))
        calinski_harabasz_str = format_metric(metrics.get('calinski_harabasz'))

        document.add_paragraph(f"Silhouette Score: {silhouette_str}")
        document.add_paragraph(f"Davies-Bouldin Index: {davies_bouldin_str}")
        document.add_paragraph(f"Calinski-Harabasz Index: {calinski_harabasz_str}")
        document.add_paragraph(
            "These metrics evaluate the quality of the clusters. "
            "Higher Silhouette and Calinski-Harabasz scores indicate better-defined clusters. "
            "Lower Davies-Bouldin scores indicate better separation between clusters."
        )

        document.add_heading('4. Data Preview (First 5 Rows)', level=2)
        table = document.add_table(rows=1, cols=data_preview_df.shape[1])
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(data_preview_df.columns):
            hdr_cells[i].text = col
        for index, row in data_preview_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

        document.add_heading('5. Cluster Visualizations', level=2)
        if pca_plot_bytes:
            document.add_paragraph("PCA Plot: Visualizes clusters in a reduced 2D space.")
            document.add_picture(io.BytesIO(pca_plot_bytes), width=Inches(6))
        else:
            document.add_paragraph("PCA plot could not be generated (e.g., less than 2 numeric features or single cluster).")

        if profile_plot_bytes:
            document.add_paragraph("Numeric Feature Mean Profiles by Cluster:")
            document.add_picture(io.BytesIO(profile_plot_bytes), width=Inches(6))
        else:
            document.add_paragraph("Numeric profile plot could not be generated.")

        document.add_heading('6. Cluster Profiles', level=2)
        
        document.add_paragraph("Average values of numeric features for each cluster:")
        if not (cluster_means_numeric.empty if cluster_means_numeric is not None else True):
            table_numeric = document.add_table(rows=1, cols=cluster_means_numeric.shape[1] + 1)
            table_numeric.style = 'Table Grid'
            hdr_cells_numeric = table_numeric.rows[0].cells
            hdr_cells_numeric[0].text = 'Cluster'
            for i, col in enumerate(cluster_means_numeric.columns):
                hdr_cells_numeric[i+1].text = col
            # Ensure cluster_id is a string when accessing with .loc
            for cluster_id, row_data in cluster_means_numeric.iterrows():
                row_cells = table_numeric.add_row().cells
                row_cells[0].text = str(cluster_id)
                for i, val in enumerate(row_data):
                    row_cells[i+1].text = f'{val:.2f}'
        else:
            document.add_paragraph("No numeric cluster means available.")

        document.add_paragraph("\nProportions of categorical feature values within each cluster:")
        if cluster_cat_proportions:
            for cat_col, proportions_df in cluster_cat_proportions.items():
                document.add_paragraph(f"  **{cat_col}:**")
                table_cat = document.add_table(rows=1, cols=proportions_df.shape[1] + 1)
                table_cat.style = 'Table Grid'
                hdr_cells_cat = table_cat.rows[0].cells
                hdr_cells_cat[0].text = 'Cluster'
                for i, col in enumerate(proportions_df.columns):
                    hdr_cells_cat[i+1].text = col
                # Ensure cluster_id is a string when accessing with .loc
                for cluster_id, row_data in proportions_df.iterrows():
                    row_cells = table_cat.add_row().cells
                    row_cells[0].text = str(cluster_id)
                    for i, val in enumerate(row_data):
                        row_cells[i+1].text = f'{val:.2%}' # Format as percentage
                document.add_paragraph() # Add space between tables
        else:
            document.add_paragraph("No categorical cluster proportions available.")

        # New Section: Cluster Summaries - Formatted as discussed
        document.add_heading('7. Cluster Summaries and Characteristics', level=2)
        document.add_paragraph(
            "Below is a detailed profile for each identified cluster, highlighting their key distinguishing features "
            "compared to the overall dataset. This can help you understand the unique characteristics of each segment "
            "and develop targeted strategies."
        )

        if not (cluster_means_numeric.empty if cluster_means_numeric is not None else True) or cluster_cat_proportions:
            # Pass labels and chosen_algo to generate_cluster_summaries for cluster size calculation and DBSCAN handling
            structured_cluster_summaries = generate_cluster_summaries(cluster_means_numeric, cluster_cat_proportions, original_df_for_profile, labels, chosen_algo)
            if structured_cluster_summaries:
                # Sort the cluster summaries by integer cluster ID for consistent order
                sorted_summary_keys = sorted(structured_cluster_summaries.keys(), key=lambda x: int(x) if x.lstrip('-').isdigit() else float('inf'))
                
                for cluster_id_key in sorted_summary_keys:
                    summary_data = structured_cluster_summaries[cluster_id_key]
                    
                    # Add Cluster Heading (bold)
                    p = document.add_paragraph()
                    run = p.add_run(summary_data["cluster_heading"])
                    run.bold = True
                    p.paragraph_format.space_after = Pt(6) # Small space after heading

                    # Add Numeric Characteristics
                    if summary_data["numeric_characteristics"]:
                        p = document.add_paragraph()
                        run = p.add_run("Numeric Characteristics:")
                        run.bold = True
                        for desc in summary_data["numeric_characteristics"]:
                            p_item = document.add_paragraph(style='List Bullet')
                            # Split to bold only the feature name
                            parts = desc.split(':', 1)
                            if len(parts) > 1:
                                run_feature = p_item.add_run(parts[0].strip() + ': ')
                                run_feature.bold = True
                                p_item.add_run(parts[1].strip())
                            else:
                                p_item.add_run(desc)
                        p.paragraph_format.space_after = Pt(6) # Small space after section

                    # Add Categorical Characteristics
                    if summary_data["categorical_characteristics"]:
                        p = document.add_paragraph()
                        run = p.add_run("Categorical Characteristics:")
                        run.bold = True
                        for desc in summary_data["categorical_characteristics"]:
                            p_item = document.add_paragraph(style='List Bullet')
                            # Split to bold only the category name (e.g., "Predominantly Female")
                            parts = desc.split(' (for ', 1)
                            if len(parts) > 1:
                                run_cat = p_item.add_run(parts[0].strip())
                                run_cat.bold = True
                                p_item.add_run(' (for ' + parts[1].strip())
                            else:
                                p_item.add_run(desc)
                        p.paragraph_format.space_after = Pt(6) # Small space after section

                    # Add Persona/Implications
                    p = document.add_paragraph()
                    run = p.add_run(summary_data["persona_implications"].replace('*', '')) # Remove asterisks as we'll italicize
                    run.italic = True
                    p.paragraph_format.space_after = Pt(12) # Larger space after each full cluster summary
            else:
                document.add_paragraph("Unable to generate detailed cluster summaries.")
        else:
            document.add_paragraph("Not enough data (numeric or categorical profiles) to generate detailed cluster summaries.")

        document.add_page_break()


    # File Upload
    st.header("1️⃣ Upload Your Data")

    st.markdown("""
    Upload your data file in **CSV or Excel** format.
    """)

    uploaded_file = st.file_uploader("Upload a CSV or Excel file", type=["csv", "xlsx"], key=f"file_uploader_{st.session_state.file_uploader_key}")

    df = None
    if uploaded_file:
        try:
            if uploaded_file.name.endswith(".csv"):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            st.success(f"✅ Successfully loaded {uploaded_file.name}")
        except Exception as e:
            st.error(f"Error loading file: {e}")

    # Data Overview and Feature Selection
    if df is not None:
        st.header("2️⃣ Data Overview")

        st.markdown("""
    Here are the first 5 rows of your data:
    """)
        st.dataframe(df.head())

        all_columns = df.columns.tolist()
        
        # Automatic ID column detection
        potential_id_cols = detect_potential_id_columns(df)
        
        st.subheader("Exclude Columns")
        st.markdown("""
        Select any columns that should **not** be used for clustering (e.g., ID columns).
        """)
        excluded_columns = st.multiselect(
            "Columns to Exclude (e.g., IDs)",
            all_columns,
            default=potential_id_cols
        )

        # Filter out excluded columns from the main selection pool
        available_columns_for_selection = [col for col in all_columns if col not in excluded_columns]

        # Pre-process columns to better infer types
        processed_numeric_cols = []
        processed_categorical_cols = []
        for col in available_columns_for_selection:
            # Try converting to numeric, handling potential non-numeric strings (like '85%')
            temp_series = pd.to_numeric(df[col], errors='coerce')
            if pd.api.types.is_numeric_dtype(temp_series) and not temp_series.isnull().all():
                processed_numeric_cols.append(col)
            elif pd.api.types.is_object_dtype(df[col]) or pd.api.types.is_categorical_dtype(df[col]):
                processed_categorical_cols.append(col)

        st.subheader("Select Numeric Features")
        st.markdown("""
    **Numeric Features:** Columns with numbers, like income, age, or transaction counts.  
    These will be standardized so that all values are on the same scale.
    """)
        selected_numeric = st.multiselect(
            "Numeric Columns",
            processed_numeric_cols,
            default=processed_numeric_cols
        )

        st.subheader("Select Categorical Features")
        st.markdown("""
    **Categorical Features:** Columns with categories, like gender, product type, or region.  
    These will be automatically converted into numeric format.
    """)
        selected_categorical = st.multiselect(
            "Categorical Columns",
            processed_categorical_cols,
            default=processed_categorical_cols
        )

        st.subheader("Handle Missing Data")
        st.markdown("""
    Choose how to handle rows with missing values in your selected columns:
    - **Drop Rows:** Remove any rows that have missing values.
    - **Impute:** Fill missing numeric values with the column average and missing categories with the most common value.
    """)
        missing_strategy = st.selectbox(
            "Missing Data Handling",
            ("drop_rows", "impute"),
            format_func=lambda x: x.replace("_", " ").title()
        )

        st.subheader("Train-Test Split")
        st.markdown("""
    Decide if you want to reserve part of your data to test how well the clustering generalizes.

    **Example:**
    - **70% Train** means 70% of the data will be used to create clusters.
    - **30% Test** will be used to validate the results.

    If you're not sure, leave it to **30%** for the test set (meaning 70% for training).
    """) # Updated text
        train_ratio = st.slider(
            "Train-Test Split Ratio",
            min_value=0.0,
            max_value=0.9,
            value=0.7, # Default changed to 0.7 (70% train, 30% test)
            step=0.1
        )

        if not selected_numeric and not selected_categorical:
            st.warning("Please select at least one numeric or categorical column for clustering.")
            st.stop()

        # Preprocess Data
        st.header("3️⃣ Data Preprocessing")

        with st.spinner("Preprocessing your data..."):
            # Ensure df passed to preprocess_data only contains selected columns
            df_for_preprocessing = df[selected_numeric + selected_categorical].copy()
            # Convert identified numeric strings to actual numbers *before* preprocessing
            for col in selected_numeric:
                df_for_preprocessing[col] = pd.to_numeric(df_for_preprocessing[col], errors='coerce')
            
            scaled_df, df_profile = preprocess_data(df_for_preprocessing, selected_numeric, selected_categorical, missing_strategy)

        st.success("✅ Preprocessing complete!")
        st.write("Here is a preview of your processed data (scaled and one-hot encoded where applicable):")
        st.dataframe(scaled_df.head())

        # Clustering Evaluation
        st.header("4️⃣ Clustering Evaluation")

        st.markdown("""
    We will now evaluate **KMeans**, **Gaussian Mixture Model**, and **Agglomerative Clustering** using different numbers of clusters.
    This helps you see which method separates your data best.
    """)

        k_range = range(2, min(11, len(scaled_df) + 1)) # Ensure k does not exceed number of samples

        with st.spinner("Evaluating clustering performance..."):
            scores, additional_metrics = evaluate_models(scaled_df, k_range)

        # Plot metrics
        fig_metrics, axes_metrics = plt.subplots(3, 1, figsize=(10, 15))
        for algo in scores:
            k_values_plot = [k_range[i] for i, score in enumerate(scores[algo]["Silhouette"]) if not np.isnan(score)]
            silhouette_scores_plot = [score for score in scores[algo]["Silhouette"] if not np.isnan(score)]
            davies_scores_plot = [score for score in scores[algo]["Davies"] if not np.isnan(score)]
            calinski_scores_plot = [score for score in scores[algo]["Calinski"] if not np.isnan(score)]

            if k_values_plot:
                axes_metrics[0].plot(k_values_plot, silhouette_scores_plot, label=algo)
                axes_metrics[1].plot(k_values_plot, davies_scores_plot, label=algo)
                axes_metrics[2].plot(k_values_plot, calinski_scores_plot, label=algo)
                
        axes_metrics[0].set_title("Silhouette Score (Higher is better)")
        axes_metrics[1].set_title("Davies-Bouldin Index (Lower is better)")
        axes_metrics[2].set_title("Calinski-Harabasz Index (Higher is better)")
        for ax in axes_metrics:
            ax.legend()
            ax.grid()
            ax.set_xlabel("Number of Clusters (k)")
        plt.tight_layout()
        st.pyplot(fig_metrics)
        plt.close(fig_metrics)

        st.markdown("""
    **How to Read These Graphs:**
    - **Silhouette Score:** Higher means clearer separation between clusters.
    - **Davies-Bouldin Index:** Lower means clusters are more compact and distinct.
    - **Calinski-Harabasz Index:** Higher means better defined clusters.
    """)

        # Plot Elbow Method (KMeans) and AIC/BIC (GMM)
        st.subheader("Elbow Method (K-Means) and Information Criteria (GMM)")
        fig_elbow, ax_elbow = plt.subplots(figsize=(10, 6))
        
        # KMeans Inertia
        k_values_kmeans = [k_range[i] for i, val in enumerate(additional_metrics["KMeans_Inertia"]) if not np.isnan(val)]
        inertia_plot = [val for val in additional_metrics["KMeans_Inertia"] if not np.isnan(val)]
        if k_values_kmeans:
            ax_elbow.plot(k_values_kmeans, inertia_plot, marker='o', label='K-Means Inertia (Lower is better)')

        # GMM AIC and BIC
        k_values_gmm = [k_range[i] for i, val in enumerate(additional_metrics["GMM_AIC"]) if not np.isnan(val)]
        aic_plot = [val for val in additional_metrics["GMM_AIC"] if not np.isnan(val)]
        bic_plot = [val for val in additional_metrics["GMM_BIC"] if not np.isnan(val)]
        if k_values_gmm:
            ax_elbow.plot(k_values_gmm, aic_plot, marker='x', linestyle='--', label='GMM AIC (Lower is better)')
            ax_elbow.plot(k_values_gmm, bic_plot, marker='s', linestyle=':', label='GMM BIC (Lower is better)')

        ax_elbow.set_title("Elbow Method for K-Means and Information Criteria for GMM")
        ax_elbow.set_xlabel("Number of Clusters (k)")
        ax_elbow.set_ylabel("Score")
        ax_elbow.legend()
        ax_elbow.grid()
        plt.tight_layout()
        st.pyplot(fig_elbow)
        plt.close(fig_elbow)

        st.markdown("""
    **How to Read Elbow/Information Criteria Plots:**
    - **K-Means Inertia (Elbow Method):** Look for a 'bend' or 'elbow' in the plot. This point often indicates a good balance between minimizing inertia and keeping the number of clusters manageable. Lower inertia is better.
    - **GMM AIC/BIC:** These criteria penalize model complexity. Lower values generally indicate a better-fitting model.
    """)

        # Model Recommendation
        st.header("🔍 Recommended Model Based on Metrics")

        recommendations = []
        for algo in scores:
            valid_silhouette_scores = [s for s in scores[algo]["Silhouette"] if not np.isnan(s)]
            if valid_silhouette_scores:
                idx_best = np.argmax(valid_silhouette_scores)
                k_best = k_range[idx_best] 
                
                silhouette_val = scores[algo]["Silhouette"][idx_best] if idx_best < len(scores[algo]["Silhouette"]) else np.nan
                davies_val = scores[algo]["Davies"][idx_best] if idx_best < len(scores[algo]["Davies"]) else np.nan
                calinski_val = scores[algo]["Calinski"][idx_best] if idx_best < len(scores[algo]["Calinski"]) else np.nan

                recommendations.append({
                    "algorithm": algo,
                    "k": k_best,
                    "silhouette": silhouette_val,
                    "davies": davies_val,
                    "calinski": calinski_val
                })

        if recommendations:
            best = max(recommendations, key=lambda x: x["silhouette"] if not np.isnan(x["silhouette"]) else -np.inf)
            
            st.success(
                f"**Recommended:** {best['algorithm']} with {best['k']} clusters "
                f"(Silhouette Score: {format_metric(best['silhouette'])}, "
                f"Davies-Bouldin: {format_metric(best['davies'])}, "
                f"Calinski-Harabasz: {format_metric(best['calinski'])}"
            )
        else:
            st.info("No clear model recommendation could be made (e.g., all scores were NaN or no valid clusters formed).")


        # Final Algorithm Selection
        st.header("5️⃣ Choose Final Model for Clustering")

        st.markdown("""
    You can choose the recommended option or pick any algorithm and parameters yourself.
    """)
        st.info("💡 To re-run clustering with different parameters on the *current dataset*, simply adjust the 'Number of clusters (k)' or 'DBSCAN' settings below and click the '🚀 Run Clustering' button again.")

        default_algo_index = 0
        if recommendations and best["algorithm"] in ["KMeans", "Gaussian Mixture Model", "Agglomerative Clustering"]:
            default_algo_index = ["KMeans", "Gaussian Mixture Model", "Agglomerative Clustering", "DBSCAN"].index(best["algorithm"])

        chosen_algo = st.selectbox(
            "Select Clustering Algorithm",
            ["KMeans", "Gaussian Mixture Model", "Agglomerative Clustering", "DBSCAN"],
            index=default_algo_index
        )

        n_clusters = None
        eps = None
        min_samples = None

        if chosen_algo in ["KMeans", "Gaussian Mixture Model", "Agglomerative Clustering"]:
            default_n_clusters = int(best["k"]) if recommendations and not np.isnan(best["k"]) else 3

            n_clusters = st.slider(
                "Number of clusters (k)",
                min_value=2,
                max_value=min(10, len(scaled_df)),
                value=default_n_clusters
            )
            if len(scaled_df) < n_clusters:
                st.warning(f"Number of clusters ({n_clusters}) exceeds the number of samples ({len(scaled_df)}). Adjust 'k'.")
                st.stop()

        else: # DBSCAN
            eps = st.slider(
                "DBSCAN: Neighborhood size (eps)",
                min_value=0.1,
                max_value=2.0,
                value=0.5,
                step=0.1
            )
            min_samples = st.slider(
                "DBSCAN: Minimum samples per cluster",
                min_value=2,
                max_value=10,
                value=5
            )

        st.markdown("""
    When you're ready, click the button below to run clustering and generate results.
    """)
        # Run Clustering
        if st.button("🚀 Run Clustering"):
            st.header("6️⃣ Clustering Results")
            with st.spinner("Running clustering..."):
                model = None
                labels = None
                silhouette, davies_bouldin, calinski_harabasz = np.nan, np.nan, np.nan

                try:
                    if chosen_algo == "KMeans":
                        model = KMeans(n_clusters=n_clusters, n_init=10, random_state=42)
                        labels = model.fit_predict(scaled_df)
                    elif chosen_algo == "Gaussian Mixture Model":
                        model = GaussianMixture(n_components=n_clusters, random_state=42)
                        labels = model.fit_predict(scaled_df)
                    elif chosen_algo == "Agglomerative Clustering":
                        model = AgglomerativeClustering(n_clusters=n_clusters)
                        labels = model.fit_predict(scaled_df)
                    elif chosen_algo == "DBSCAN":
                        model = DBSCAN(eps=eps, min_samples=min_samples)
                        labels = model.fit_predict(scaled_df)
                    
                    # Exclude DBSCAN noise points (-1) from metric calculation if they exist
                    if chosen_algo == "DBSCAN" and -1 in np.unique(labels):
                        # Filter out noise points for metric calculation
                        core_samples_mask = labels != -1
                        labels_filtered = labels[core_samples_mask]
                        scaled_df_filtered = scaled_df[core_samples_mask]
                        if len(np.unique(labels_filtered)) > 1:
                            silhouette = silhouette_score(scaled_df_filtered, labels_filtered)
                            davies_bouldin = davies_bouldin_score(scaled_df_filtered, labels_filtered)
                            calinski_harabasz = calinski_harabasz_score(scaled_df_filtered, labels_filtered)
                        elif len(np.unique(labels_filtered)) == 1:
                             st.warning(f"Only 1 cluster formed after excluding noise points. Metrics will be N/A.")
                        else: # No clusters formed after excluding noise
                            st.warning(f"No clusters formed after excluding noise points. Metrics will be N/A.")
                    elif labels is not None and len(np.unique(labels)) > 1:
                        silhouette = silhouette_score(scaled_df, labels)
                        davies_bouldin = davies_bouldin_score(scaled_df, labels)
                        calinski_harabasz = calinski_harabasz_score(scaled_df, labels)
                    elif labels is not None:
                        st.warning(f"Only {len(np.unique(labels))} cluster(s) formed. Metrics will be N/A.")

                except Exception as e:
                    st.error(f"Error during clustering: {e}. Please check your data and parameters.")
                    labels = None

            if labels is not None and (len(np.unique(labels)) > 1 or (chosen_algo == "DBSCAN" and -1 in np.unique(labels))):
                st.session_state.analysis_completed = True
                if chosen_algo == "DBSCAN":
                    num_clusters_dbscan = len(np.unique(labels))
                    if -1 in np.unique(labels):
                        num_clusters_dbscan -= 1 # Exclude noise points
                        st.success(f"✅ Clustering completed. Found {num_clusters_dbscan} clusters and {np.sum(labels == -1)} noise points.")
                    else:
                        st.success(f"✅ Clustering completed. Found {num_clusters_dbscan} clusters.")
                else:
                    st.success(f"✅ Clustering completed. Found {len(np.unique(labels))} clusters.")

                cluster_counts_df = pd.DataFrame({"Cluster": np.unique(labels), "Count": np.bincount(labels[labels != -1]) if -1 in labels else np.bincount(labels)})
                if chosen_algo == "DBSCAN" and -1 in np.unique(labels):
                    cluster_counts_df.loc[cluster_counts_df["Cluster"] == -1, "Cluster"] = "Noise (-1)"
                st.write("**Cluster Distribution:**")
                st.dataframe(cluster_counts_df)

                st.write(f"**Silhouette Score:** {format_metric(silhouette)}")
                st.write(f"**Davies-Bouldin Index:** {format_metric(davies_bouldin)}")
                st.write(f"**Calinski-Harabasz Index:** {format_metric(calinski_harabasz)}")

                # Ensure df_profile has the cluster labels for accurate plot and summary generation
                df_profile_with_labels = df_profile.copy()
                df_profile_with_labels['Cluster'] = labels # Keep labels as integers here

                fig_pca, fig_profile_numeric, cluster_means_numeric, cluster_cat_proportions = generate_plots(
                    df_profile_with_labels, labels, selected_numeric
                )

                col1, col2 = st.columns(2)
                with col1:
                    if fig_pca:
                        st.pyplot(fig_pca)
                    else:
                        st.info("PCA plot could not be generated (e.g., less than 2 numeric features or single cluster).")
                with col2:
                    if fig_profile_numeric:
                        st.pyplot(fig_profile_numeric)
                    else:
                        st.info("Numeric cluster profile plot could not be generated.")

                st.subheader("Numeric Feature Means per Cluster")
                if not (cluster_means_numeric.empty if cluster_means_numeric is not None else True):
                    st.dataframe(cluster_means_numeric.round(2))
                else:
                    st.info("No numeric cluster means to display (check selected features).")

                if selected_categorical:
                    st.subheader("Categorical Feature Distributions per Cluster")
                    for cat in selected_categorical:
                        # Use df_profile_with_labels for calculating proportions to ensure consistency
                        temp_df_for_cat_prop = df_profile_with_labels[df_profile_with_labels['Cluster'] != -1] if chosen_algo == "DBSCAN" else df_profile_with_labels
                        
                        # Only calculate if there are actual clusters remaining after filtering (if any)
                        if not temp_df_for_cat_prop.empty and len(np.unique(temp_df_for_cat_prop['Cluster'])) > 0:
                            prop_df = pd.crosstab(temp_df_for_cat_prop["Cluster"], temp_df_for_cat_prop[cat], normalize="index")
                            st.markdown(f"**{cat}:**")
                            st.dataframe((prop_df * 100).round(1))
                        else:
                            st.info(f"No valid clusters to display categorical proportions for '{cat}' after filtering noise (if any).")
                else:
                    st.info("No categorical cluster proportions to display (no categorical features selected).")

                st.subheader("7️⃣ Download Results")
                st.info("You can download the clustered data or a full report summarizing the analysis.")

                document = Document()
                pca_plot_bytes = io.BytesIO()
                if fig_pca:
                    fig_pca.savefig(pca_plot_bytes, format='png', bbox_inches='tight')
                    pca_plot_bytes.seek(0)
                else:
                    pca_plot_bytes = None

                profile_plot_bytes = io.BytesIO()
                if fig_profile_numeric:
                    fig_profile_numeric.savefig(profile_plot_bytes, format='png', bbox_inches='tight')
                    profile_plot_bytes.seek(0)
                else:
                    profile_plot_bytes = None

                report_bytes_io = io.BytesIO()
                create_report(
                    document,
                    chosen_algo,
                    {'n_clusters': n_clusters, 'eps': eps, 'min_samples': min_samples} if chosen_algo != 'DBSCAN' else {'eps': eps, 'min_samples': min_samples},
                    {'silhouette': silhouette, 'davies_bouldin': davies_bouldin, 'calinski_harabasz': calinski_harabasz},
                    df_profile_with_labels.head(5), # Use df_profile_with_labels for preview
                    pca_plot_bytes.getvalue() if pca_plot_bytes else None,
                    profile_plot_bytes.getvalue() if profile_plot_bytes else None,
                    cluster_means_numeric,
                    cluster_cat_proportions,
                    df_profile, # Pass original df_profile without labels for overall stats
                    labels, # Pass labels here for cluster size calculation in report
                    chosen_algo # Pass chosen_algo for DBSCAN noise handling
                )
                document.save(report_bytes_io)
                report_bytes = report_bytes_io.getvalue()
                report_bytes_io.close()

                st.download_button(
                    label="📥 Download Comprehensive Report (.docx)",
                    data=report_bytes,
                    file_name=f"ML_Analysis_Report_{pd.to_datetime('today').strftime('%Y%m%d_%H%M%S')}.docx", # Report filename updated
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                csv_buffer = io.StringIO()
                df_profile_with_labels.to_csv(csv_buffer, index=False) # Save df_profile_with_labels for download
                st.download_button(
                    label="📥 Download Clustered Data (CSV)",
                    data=csv_buffer.getvalue(),
                    file_name=f"Clustered_Customer_Data_{pd.to_datetime('today').strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                )
            else:
                st.error("Clustering did not form more than one cluster or failed. Please review your data and selected parameters.")
                st.session_state.analysis_completed = False

    # Reset Buttons and "What's Next" section
    if st.session_state.analysis_completed:
        st.header("🎯 Analysis Complete")
        st.markdown("You can either re-run clustering with different parameters on the current dataset, or clear everything to start fresh with new data.")
        col_reset1, col_reset2 = st.columns(2)
        with col_reset1:
            if st.button("🔄 Rerun Clustering with New Parameters"):
                st.session_state.analysis_completed = False
                st.rerun()
        with col_reset2:
            if st.button("🗑️ Clear All Data & Start Fresh"):
                # Preserve 'authenticated' state, clear others
                auth_status = st.session_state.get('authenticated', False)
                
                # Increment file_uploader_key for a fresh file uploader widget
                current_file_uploader_key = st.session_state.get('file_uploader_key', 0)
                
                # Clear all session state items EXCEPT 'authenticated'
                for key in list(st.session_state.keys()):
                    if key != 'authenticated':
                        del st.session_state[key]
                
                # Re-set 'authenticated' and update file_uploader_key
                st.session_state.authenticated = auth_status
                st.session_state.file_uploader_key = current_file_uploader_key + 1
                
                st.rerun()
    else:
        if df is None:
            st.info("Please upload a data file (.csv or .xlsx) at the top of the page to begin the analysis.")

    # --- Copyright Notice for Main App ---
    st.markdown("---") # Optional: add a separator line
    st.markdown("<p style='text-align: center; color: grey;'>© Copyright SSO Consultants</p>", unsafe_allow_html=True)


# --- Run the appropriate page based on authentication status ---
if not st.session_state.authenticated:
    login_page()
else:
    main_app()
